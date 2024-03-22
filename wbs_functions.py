from docx import Document
from datetime import datetime


#To-Be Tables List
tables_list_info = [
    (0, 'Revision history'),
    (1, 'Process Map'),
    (2, 'Performers'),
    (3, 'Activities'),
    (4, 'DOTMIL'),
    (5, 'Pain Points'),
    (6, 'Systems'),
    (7, 'FG MAP'),
    (8, 'FG results'),
    (9, 'KPI')
]

# List of LOOs and the Pictures file names
LOO_List = [
    {'Excel_File_Name': 'Commercial-Line-Haul', 'Normal_Map': 'Commercial-Line-Haul_Pic', 'FG_Map': 'Commercial-Line-Haul_FG'},
    {'Excel_File_Name': 'Carrier_Function', 'Normal_Map': 'Carrier_Function_Pic', 'FG_Map': 'Carrier_Function_FG'},
    {'Excel_File_Name': 'Rate_Management', 'Normal_Map': 'Rate_Management_Pic', 'FG_Map': 'Rate_Management_FG'},
    {'Excel_File_Name': 'Payment_Process', 'Normal_Map': 'Payment_Process_Pic', 'FG_Map': 'Payment_Process_FG'},
    {'Excel_File_Name': 'Deployment', 'Normal_Map': 'Deployment_Pic', 'FG_Map': 'Deployment_FG'},
    {'Excel_File_Name': 'Rail', 'Normal_Map': 'Rail_Pic', 'FG_Map': 'Rail_FG'},
    {'Excel_File_Name': 'Inbound_Terminating_Freight', 'Normal_Map': 'Inbound_Terminating_Freight_Pic', 'FG_Map': 'Inbound_Terminating_Freight_FG'},
    {'Excel_File_Name': 'Support_HQ', 'Normal_Map': 'Support_HQ_Pic', 'FG_Map': 'Support_HQ_FG'},
    {'Excel_File_Name': 'Outbound', 'Normal_Map': 'Outbound_Pic', 'FG_Map': 'Outbound_FG'},
    {'Excel_File_Name': 'HAZMAT', 'Normal_Map': 'HAZMAT_Pic', 'FG_Map': 'HAZMAT_FG'},
    {'Excel_File_Name': 'TPS', 'Normal_Map': 'TPS_Pic', 'FG_Map': 'TPS_FG'},
    {'Excel_File_Name': 'ITV', 'Normal_Map': 'ITV_Pic', 'FG_Map': 'ITV_FG'},
    {'Excel_File_Name': 'TAC', 'Normal_Map': 'TAC_Pic', 'FG_Map': 'TAC_FG'}
]

# Placeholders names in the word doc
placeholders_list = [
    'REVISION__TABLE',
    'PERFORMERS__TABLE',
    'ACTIVITIES__TABLE',
    'DOTMIL__TABLE',
    'FINDINGS__TABLE',
    'SYSTEMS__TABLE',
    'FITGAP__TABLE',
    'KPI__TABLE'
]


# function to move the table after the placeholder
def move_table (document, table, placeholder):
    # Find the index of the paragraph containing the text "Coverpage"
    target_text = placeholder
    
    target_paragraph_index = None
    for i, paragraph in enumerate(document.paragraphs):
        if target_text in paragraph.text:
            target_paragraph_index = i
            break
    # Moving table Function
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)
    
    # Insert the table after the paragraph containing "Coverpage"
    if target_paragraph_index is not None:
        paragraph = document.paragraphs[target_paragraph_index]
        move_table_after(table, paragraph)
    
    #Clear the Placeholder
    paragraph.clear()


# Function to set the tables to autofit
def set_autofit(doc: Document):
    """
    Hotfix for autofit.
    """
    for t_idx, table in enumerate(doc.tables):
        doc.tables[t_idx].autofit = True
        doc.tables[t_idx].allow_autofit = True
        doc.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
        for row_idx, r_val in enumerate(doc.tables[t_idx].rows):
            for cell_idx, c_val in enumerate(doc.tables[t_idx].rows[row_idx].cells):
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
                doc.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0
    return doc


# Function to fix date formating in the revision table (table[0])
def fix_date(document):
    table = document.tables[0]
    
    # Iterate through the second column of the first table and reformat dates
    for row in table.rows:
        cell = row.cells[1]
        # Assuming the date string is in the format "YYYY-MM-DD HH:MM:SS"
        date_string = cell.text.strip()
        try:
            # Parse the date string into a datetime object
            date_object = datetime.strptime(date_string, '%Y-%m-%d %H:%M:%S')
            # Format the datetime object into "MM/DD/YYYY" format
            formatted_date = date_object.strftime('%m/%d/%Y')
            # Update the cell text with the formatted date
            cell.text = formatted_date
            
            # Set style "Normal1" to the paragraph in the cell
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal1"
                
        except ValueError:
            # Handle any potential errors in date string parsing
            # print(f"Error parsing date in cell: {cell.text}")
            pass

# Function to delete a column from a table (See table list for index values)
def Delete_column_in_table(table, columns, document):
    table = document.tables[table]
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for cell in table.column_cells(columns):
        cell._tc.getparent().remove(cell._tc)
    col_elem = grid[columns]
    grid.remove(col_elem)