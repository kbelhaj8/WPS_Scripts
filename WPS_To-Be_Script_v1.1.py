import os
import openpyxl
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


#import Move_Table
import wbs_functions

# Get the current working directory
folder_path = os.path.join(os.getcwd(), "Excel_Reports")

# Load the Word document
document = Document('wps_script_template.docx')

# Iterate through Excel files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith('.xlsx'):
        workbook = openpyxl.load_workbook(os.path.join(folder_path, filename))

        # list the spreadsheet's tabs and then remove the first two (Cover page and doc info)
        tabs_list = workbook.sheetnames
        tabs_list.remove('CoverSheet')
        tabs_list.remove('ProcessInfo')
        tabs_list.remove('Process Overview')
        
        
        #@@@@Copy All the tables@@@@
        # Iterate through each tab in tabs_list and its corresponding placeholder in placeholders_list
        for tab, placeholder_name in zip(tabs_list, wbs_functions.placeholders_list):
            sheet = workbook[tab]
            table_data = []
            for row in sheet.iter_rows(values_only=True):
                table_data.append(row)
        
            # Add table to the Word document
            table = document.add_table(rows=0, cols=len(table_data[0]))
            table.style = 'WWC_Table1'
            
            # Add data to the table
            for row_data in table_data:
                row_cells = table.add_row().cells
                for j, value in enumerate(row_data):
                    cell = row_cells[j]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        paragraph.style = "Normal1"  # Set the style to "Normal1" for each paragraph
        
        
            # Move the table to the desired placeholder
            wbs_functions.move_table(document, table, placeholder_name)
        
        #@@@Update The Process Overview (Paragraph)
        # Retrieve data from the "Process Overview" tab for the Process Summary section
        process_overview_sheet = workbook['Process Overview']
        process_summary_data = process_overview_sheet['A2'].value  # text in cell A2
        
        # Find and replace the placeholder with the process summary
        for paragraph in document.paragraphs:
            if 'PROCESS__SUMMARY' in paragraph.text:
                paragraph.text = process_summary_data
                break
        
        #@@@Update The Cover Page with LOO name and Date
        # Update the LOO Name in the Cover Page
        Process_Info_sheet = workbook['ProcessInfo']
        LOO_name_data = Process_Info_sheet['B3'].value  # text in cell B3
        
        # Find and replace the placeholder with the process summary
        for paragraph in document.paragraphs:
            if 'LOO__NAME' in paragraph.text:
                paragraph.text = LOO_name_data
                paragraph.style = "Cover_1"
                break
        
        # Update the Date in the Cover Page
        Process_Info_sheet = workbook['ProcessInfo']
        LOO_date_data = Process_Info_sheet['B4'].value  # text in cell B4
        
        # Format the datetime object as a string
        formatted_date = LOO_date_data.strftime("%m/%d/%Y")  # Change the format as needed
        
        # Find and replace the placeholder with the process summary
        for paragraph in document.paragraphs:
            if 'LOO__Date' in paragraph.text:
                paragraph.text = formatted_date  # Assign the formatted date string
                paragraph.style = "Cover_1"
                break
        
        #@@@ Add Process maps to the document inside tables
        tables = document.tables

        # Set the folder path to "Visio Screenshots"
        visio_folder_path = os.path.join(os.getcwd(), "Visio_Screenshots")

        # Iterate through LOO_List and add pictures to the tables
        for loo_item in wbs_functions.LOO_List:
            if loo_item['Excel_File_Name'] in filename:
                # Create a blank table and add the picture (non FITGAP)
                p1 = tables[1].rows[0].cells[0].add_paragraph()
                p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align the paragraph to the middle
                r1 = p1.add_run()
                visio_screenshot_path = os.path.join(visio_folder_path, f"{loo_item['Normal_Map']}.png")
                r1.add_picture(visio_screenshot_path, width=Inches(8.0))
        
                # Create a blank table and add the picture (FITGAP)
                p2 = tables[7].rows[0].cells[0].add_paragraph()
                p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align the paragraph to the middle
                r2 = p2.add_run()
                visio_screenshot_path_fg = os.path.join(visio_folder_path, f"{loo_item['FG_Map']}.png")
                r2.add_picture(visio_screenshot_path_fg, width=Inches(8.0))
                break


        # Call function to fix the date format
        wbs_functions.fix_date(document)
        
        # Delete the first two columns from the activities table
        wbs_functions.Delete_column_in_table(3, 0, document)
        wbs_functions.Delete_column_in_table(3, 0, document)
        
        # Delete the extra columns from the DOTMIL table
        wbs_functions.Delete_column_in_table(4, 0, document)
        wbs_functions.Delete_column_in_table(4, 0, document)
        wbs_functions.Delete_column_in_table(4, 0, document)
        wbs_functions.Delete_column_in_table(4, 5, document)
        
        # Delete the extra columns from the PP table
        wbs_functions.Delete_column_in_table(5, 0, document)
        wbs_functions.Delete_column_in_table(5, 2, document)
        wbs_functions.Delete_column_in_table(5, 3, document)
        
        # Delete the extra columns from the FG table
        wbs_functions.Delete_column_in_table(8, 0, document)
        wbs_functions.Delete_column_in_table(8, 0, document)
        
        
        #Call function to resize the tables to autofit
        wbs_functions.set_autofit(document)


        ###@@@ Output results into a seperate folder
        # Set the target folder path to save the documents
        output_folder_path = os.path.join(os.getcwd(), "WPS_Docs")

        # Create the output folder if it doesn't exist
        if not os.path.exists(output_folder_path):
            os.makedirs(output_folder_path)
    
        # Save the updated document to the output folder
        output_filename = os.path.splitext(filename)[0].replace("Reports", "WPS") + '.docx'
        output_path = os.path.join(output_folder_path, output_filename)
        document.save(output_path)
        print(f"Document '{output_filename}' has been created successfully in 'WPS_Docs' folder.")
        
        # Reset the document for the next iteration
        document = Document('wps_script_template.docx')
